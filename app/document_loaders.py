from pathlib import Path


def load_document(path: str) -> str:
    suffix = Path(path).suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in {".docx", ".doc"}:
        return _load_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _load_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf. Run: pip install -r requirements.txt") from exc
    reader = PdfReader(path)
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {index}]\n{text}")
    return "\n\n".join(pages)


def _load_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Word support requires python-docx. Run: pip install -r requirements.txt") from exc
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
